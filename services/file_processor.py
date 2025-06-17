import os
import PyPDF2
from PIL import Image
from docx import Document
import logging

# OCR is optional - graceful fallback if not available
try:
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

class FileProcessor:
    def __init__(self):
        self.supported_formats = ['pdf', 'txt', 'png', 'jpg', 'jpeg', 'gif', 'docx']
    
    def extract_text(self, file_path, file_type):
        """Extract text from uploaded file based on file type"""
        try:
            if file_type.lower() == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type.lower() == 'txt':
                return self._extract_from_txt(file_path)
            elif file_type.lower() in ['png', 'jpg', 'jpeg', 'gif']:
                return self._extract_from_image(file_path)
            elif file_type.lower() == 'docx':
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logging.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logging.error(f"Error reading PDF: {str(e)}")
            return None
    
    def _extract_from_txt(self, file_path):
        """Extract text from text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logging.error(f"Error reading text file: {str(e)}")
                return None
        except Exception as e:
            logging.error(f"Error reading text file: {str(e)}")
            return None
    
    def _extract_from_image(self, file_path):
        """Extract text from image using OCR"""
        try:
            # Open image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Use OCR to extract text if available
            if OCR_AVAILABLE:
                try:
                    text = pytesseract.image_to_string(image, lang='eng')
                    return text.strip()
                except Exception as ocr_error:
                    logging.warning(f"OCR failed: {str(ocr_error)}. Returning image metadata.")
            
            # Return basic image info if OCR not available or fails
            return f"Image file: {os.path.basename(file_path)}\nSize: {image.size}\nMode: {image.mode}\nOCR processing not available - please ensure text is clearly visible in the image."
                
        except Exception as e:
            logging.error(f"Error processing image: {str(e)}")
            return None
    
    def _extract_from_docx(self, file_path):
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logging.error(f"Error reading DOCX file: {str(e)}")
            return None
    
    def validate_file(self, file_path, file_type):
        """Validate if file is readable and contains content"""
        if not os.path.exists(file_path):
            return False, "File not found"
        
        if file_type.lower() not in self.supported_formats:
            return False, f"Unsupported file type: {file_type}"
        
        try:
            content = self.extract_text(file_path, file_type)
            if not content or len(content.strip()) < 10:
                return False, "File appears to be empty or contains insufficient text"
            return True, "File is valid"
        except Exception as e:
            return False, f"Error validating file: {str(e)}"
    
    def get_file_info(self, file_path):
        """Get basic information about the file"""
        try:
            stat = os.stat(file_path)
            return {
                'size': stat.st_size,
                'modified': stat.st_mtime,
                'readable': os.access(file_path, os.R_OK)
            }
        except Exception as e:
            logging.error(f"Error getting file info: {str(e)}")
            return None
